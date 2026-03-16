"""
build_chro_approval.py - Build the CHRO pre-load approval document as a PDF.

Reads existing run artifacts only:
  - audit_trail.json
  - sanity_gate.json
  - corrections_manifest.csv
  - config/policy.yaml

The script renders a temporary DOCX and converts it to:
  dashboard_runs/{run_id}/chro_approval_document.pdf
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import tempfile
import textwrap
from datetime import datetime
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", str(Path(tempfile.gettempdir()) / "recon_kit_mpl"))
os.environ.setdefault("MPLBACKEND", "Agg")

import matplotlib.pyplot as plt
import pandas as pd
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.patches import Rectangle

plt.rcParams["pdf.compression"] = 0

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("[error] python-docx not installed", file=sys.stderr)
    sys.exit(2)

HERE = Path(__file__).resolve().parent
ROOT = HERE.parents[1]
sys.path.insert(0, str(ROOT / "audit" / "summary"))

from config_loader import load_policy
BLUE_DARK = RGBColor(0x1F, 0x4E, 0x79)
GREY_FILL = "F3F6F9"
WARN_FILL = "FDE9D9"
BORDER = "7F8C8D"


def _add_page_number_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def _set_cell_border(cell, color: str = BORDER, size: str = "12") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _style_doc(doc: Document, run_id: str, report_date: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = BLUE_DARK
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = BLUE_DARK

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run(
        f"CONFIDENTIAL - Data Whisperer Pre-Load Approval Document "
        f"Run {run_id} - Generated {report_date} - Page "
    )
    _add_page_number_field(footer_p, "PAGE")
    footer_p.add_run(" of ")
    _add_page_number_field(footer_p, "NUMPAGES")
    for run in footer_p.runs:
        run.font.size = Pt(8)


def _load_json(path: Path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def _load_manifest(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame(
            columns=[
                "correction_type",
                "worker_id",
                "pair_id",
                "match_source",
                "fix_types",
                "action",
                "confidence",
                "summary",
                "output_file",
            ]
        )
    return pd.read_csv(path)


def _count_reasons(decisions: list[dict], needle: str) -> int:
    needle = needle.lower()
    count = 0
    for row in decisions:
        reason = str(row.get("reason") or "").lower()
        if needle in reason:
            count += 1
    return count


def _pct(part: int, total: int) -> str:
    return f"{((part / total) * 100) if total else 0:.1f}"


def _read_run_data(run_dir: Path) -> dict:
    audit_trail = _load_json(run_dir / "audit_trail.json", {})
    sanity_gate = _load_json(run_dir / "sanity_gate.json", {})
    manifest = _load_manifest(run_dir / "corrections_manifest.csv")
    policy = load_policy(ROOT / "config" / "policy.yaml")

    decisions = audit_trail.get("action_decisions", []) or []
    total_records = int(audit_trail.get("total_records_processed", 0) or 0)
    approve_count = int(audit_trail.get("approve_count", 0) or 0)
    review_count = int(audit_trail.get("review_count", 0) or 0)
    reject_count = int(audit_trail.get("reject_count", 0) or 0)
    corrections_staged = int(audit_trail.get("corrections_staged_count", 0) or 0)

    correction_counts = (
        manifest["correction_type"].value_counts().to_dict()
        if "correction_type" in manifest.columns else {}
    )

    gate_metrics = sanity_gate.get("metrics", {}) or {}
    gate_health = sanity_gate.get("health_checks", {}) or {}
    active_zero = int(
        ((gate_health.get("active_zero_salary") or {}).get("value", 0)) or 0
    )
    hire_date_wave = int(
        ((gate_metrics.get("hire_date_wave") or {}).get("count", 0)) or 0
    ) or _count_reasons(decisions, "hire_date_wave")
    salary_variance = _count_reasons(decisions, "salary_ratio_extreme")

    return {
        "policy": policy,
        "audit_trail": audit_trail,
        "sanity_gate": sanity_gate,
        "manifest": manifest,
        "client_name": str(policy.get("client_name") or "Your Organization"),
        "total_records": total_records,
        "approve_count": approve_count,
        "review_count": review_count,
        "reject_count": reject_count,
        "approve_pct": _pct(approve_count, total_records),
        "review_pct": _pct(review_count, total_records),
        "job_org_count": int(correction_counts.get("job_org", 0) or 0),
        "hire_date_count": int(correction_counts.get("hire_date", 0) or 0),
        "status_count": int(correction_counts.get("status", 0) or 0),
        "salary_count": int(correction_counts.get("salary", 0) or 0),
        "zero_salary_count": active_zero,
        "hire_date_wave_count": hire_date_wave,
        "salary_variance_count": salary_variance,
        "gate_result": "PASS" if sanity_gate.get("passed", False) else "FAIL",
        "gate_reasons": sanity_gate.get("reasons", []) or [],
        "corrections_staged_count": corrections_staged,
    }


def _add_title(doc: Document, text: str, size: int = 24) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK


def _add_cover_page(doc: Document, run_id: str, data: dict, report_date: str) -> None:
    _add_title(doc, "DATA WHISPERER", 24)
    _add_title(doc, "Pre-Load Approval Document", 18)
    doc.add_paragraph()

    rows = [
        ("Organization:", data["client_name"]),
        ("Migration:", "ADP to Workday"),
        ("Run ID:", run_id),
        ("Report Date:", report_date),
        ("Prepared By:", "Data Whisperer Reconciliation Engine"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        left.width = Inches(2.3)
        right.width = Inches(4.5)
        left.text = label
        right.text = value
        _set_cell_fill(left, GREY_FILL)
        _set_cell_border(left)
        _set_cell_border(right)
        for run in left.paragraphs[0].runs:
            run.font.bold = True

    doc.add_paragraph()
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIAL - FOR EXECUTIVE REVIEW ONLY")
    run.font.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    sig_table = doc.add_table(rows=1, cols=1)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_cell = sig_table.cell(0, 0)
    _set_cell_border(sig_cell, color="1F4E79")
    sig_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = sig_cell.paragraphs[0]
    p.add_run("CHIEF HUMAN RESOURCES OFFICER APPROVAL").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "",
        "[ ] I have reviewed this reconciliation summary and authorize",
        "    the correction load described in this document.",
        "",
        "Signature: _______________________________",
        "Name (Print): ___________________________",
        "Title: __________________________________",
        "Date: ___________________________________",
    ]:
        para = sig_cell.add_paragraph(line)
        if line.startswith("[ ]"):
            para.runs[0].font.bold = True


def _add_executive_summary(doc: Document, data: dict) -> None:
    doc.add_heading("What We Found", level=1)
    doc.add_paragraph(
        f"We compared {data['total_records']:,} employee records between your ADP "
        f"system and Workday. Of those, {data['approve_count']:,} records "
        f"({data['approve_pct']}%) matched cleanly and are ready to load. "
        f"{data['review_count']:,} records ({data['review_pct']}%) require human review "
        f"before they can be loaded. {data['reject_count']:,} records were flagged "
        f"as possible wrong-person matches and have been blocked entirely."
    )
    doc.add_paragraph(
        f"If you authorize this load, the following corrections will be applied to "
        f"{data['approve_count']:,} Workday records:\n"
        f"- {data['job_org_count']:,} job title or department corrections\n"
        f"- {data['hire_date_count']:,} hire date corrections\n"
        f"- {data['status_count']:,} employment status corrections\n"
        f"- {data['salary_count']:,} salary corrections\n"
        f"No corrections will be applied to records in the review queue until a human reviewer clears them."
    )
    doc.add_paragraph(
        f"The following records were automatically protected and will NOT receive any corrections "
        f"regardless of authorization:\n"
        f"- {data['zero_salary_count']:,} active employees showing $0 salary in Workday "
        f"(possible data entry error - requires investigation)\n"
        f"- {data['review_count']:,} records flagged for human review\n"
        f"- {data['reject_count']:,} records blocked as possible wrong-person matches"
    )


def _risk_rows(data: dict) -> list[tuple[str, int, str]]:
    rows = [
        ("Active employees with $0 salary", data["zero_salary_count"], "Investigate before load"),
        ("Records requiring human review", data["review_count"], "Complete review queue first"),
        ("Possible wrong-person matches", data["reject_count"], "Manual investigation required"),
        ("Hire date wave detected", data["hire_date_wave_count"], "Verify bulk import date is correct"),
        ("Salary changes > 15% variance", data["salary_variance_count"], "Spot check recommended"),
    ]
    return [row for row in rows if row[1] > 0]


def _add_risk_summary(doc: Document, data: dict) -> None:
    doc.add_heading("Risks and Flags Identified", level=1)
    rows = _risk_rows(data)
    if rows:
        table = doc.add_table(rows=1 + len(rows), cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Risk Type", "Count", "Action Required"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            _set_cell_fill(cell, GREY_FILL)
            _set_cell_border(cell)
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, value in enumerate(row):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = f"{value:,}" if isinstance(value, int) else str(value)
                _set_cell_border(cell)
    else:
        doc.add_paragraph("No material risk rows were identified for this run.")

    doc.add_paragraph()
    doc.add_paragraph(
        f"The sanity gate result for this run was {data['gate_result']}. "
        f"If the gate result is FAIL, corrections have been blocked and this document "
        f"should not be signed until the gate failure reason has been investigated and resolved."
    )

    if data["gate_result"] == "FAIL":
        warn_table = doc.add_table(rows=1, cols=1)
        warn_cell = warn_table.cell(0, 0)
        _set_cell_border(warn_cell, color="C0504D", size="16")
        _set_cell_fill(warn_cell, WARN_FILL)
        warn_cell.text = (
            "WARNING: SANITY GATE FAILED\n"
            "Corrections are currently blocked. Do not sign this document\n"
            "until the gate failure has been reviewed and resolved.\n"
            f"Gate failure reason: {'; '.join(data['gate_reasons']) if data['gate_reasons'] else 'No reason provided'}"
        )
        for para in warn_cell.paragraphs:
            for run in para.runs:
                run.font.bold = True


def _add_authorization(doc: Document, run_id: str, data: dict) -> None:
    audit_trail = data["audit_trail"]
    input_files = audit_trail.get("input_files", {}) or {}
    old_file = input_files.get("old_system", {}) or {}
    new_file = input_files.get("new_system", {}) or {}

    doc.add_heading("Authorization Record", level=1)
    rows = [
        ("Run ID:", run_id),
        ("Run completed:", str(audit_trail.get("run_complete_timestamp") or "")),
        ("Engine version:", str(audit_trail.get("engine_version") or "")),
        (
            "Old system file:",
            f"{old_file.get('filename', '')} (SHA-256: {str(old_file.get('sha256') or '')[:16]}...)",
        ),
        (
            "New system file:",
            f"{new_file.get('filename', '')} (SHA-256: {str(new_file.get('sha256') or '')[:16]}...)",
        ),
        ("Total records processed:", f"{data['total_records']:,}"),
        ("Records approved:", f"{data['approve_count']:,}"),
        ("Records for review:", f"{data['review_count']:,}"),
        ("Records blocked:", f"{data['reject_count']:,}"),
        ("Corrections staged:", f"{data['corrections_staged_count']:,}"),
        ("Gate result:", data["gate_result"]),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        left.text = label
        right.text = value
        _set_cell_fill(left, GREY_FILL)
        _set_cell_border(left)
        _set_cell_border(right)
        for run in left.paragraphs[0].runs:
            run.font.bold = True

    doc.add_heading("What This Signature Authorizes", level=2)
    doc.add_paragraph(
        f"By signing this document, the Chief Human Resources Officer confirms they have reviewed "
        f"the reconciliation summary above and authorizes Data Whisperer to apply the staged "
        f"corrections to {data['approve_count']:,} Workday employee records. This authorization "
        f"does not apply to records in the review queue or to records that have been blocked. "
        f"Those records require separate disposition before any corrections can be applied."
    )
    doc.add_heading("Document Retention", level=2)
    doc.add_paragraph(
        "This signed document should be retained in accordance with your organization's HR records "
        "retention policy. The run ID above can be used to retrieve the full technical audit trail at any time."
    )


def _pdf_page_base(fig, run_id: str, report_date: str, page_num: int, total_pages: int) -> None:
    fig.text(
        0.5,
        0.025,
        f"CONFIDENTIAL - Data Whisperer Pre-Load Approval Document "
        f"Run {run_id} - Generated {report_date} - Page {page_num} of {total_pages}",
        ha="center",
        va="center",
        fontsize=8,
    )


def _pdf_wrapped(fig, x: float, y: float, text: str, width: int = 86, fontsize: int = 11, weight: str = "normal") -> None:
    fig.text(x, y, textwrap.fill(text, width=width), ha="left", va="top", fontsize=fontsize, weight=weight)


def _build_pdf(run_id: str, data: dict, out_path: Path) -> None:
    report_date = datetime.now().strftime("%B %d, %Y")
    total_pages = 4
    out_path.parent.mkdir(parents=True, exist_ok=True)

    with PdfPages(out_path) as pdf:
        # Page 1
        fig = plt.figure(figsize=(8.5, 11))
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        fig.text(0.5, 0.91, "DATA WHISPERER", ha="center", va="center", fontsize=24, weight="bold", color="#1F4E79")
        fig.text(0.5, 0.87, "Pre-Load Approval Document", ha="center", va="center", fontsize=18, weight="bold", color="#1F4E79")
        info_y = 0.78
        info_lines = [
            f"Organization: {data['client_name']}",
            "Migration: ADP to Workday",
            f"Run ID: {run_id}",
            f"Report Date: {report_date}",
            "Prepared By: Data Whisperer Reconciliation Engine",
        ]
        for idx, line in enumerate(info_lines):
            fig.text(0.18, info_y - (idx * 0.04), line, ha="left", va="center", fontsize=12)
        fig.text(0.5, 0.56, "CONFIDENTIAL - FOR EXECUTIVE REVIEW ONLY", ha="center", va="center", fontsize=11, weight="bold")
        ax.add_patch(Rectangle((0.12, 0.10), 0.76, 0.30, fill=False, linewidth=1.8, edgecolor="#1F4E79"))
        fig.text(0.5, 0.37, "CHIEF HUMAN RESOURCES OFFICER APPROVAL", ha="center", va="center", fontsize=14, weight="bold")
        approval_lines = [
            "[ ] I have reviewed this reconciliation summary and authorize",
            "    the correction load described in this document.",
            "",
            "Signature: _______________________________",
            "Name (Print): ___________________________",
            "Title: __________________________________",
            "Date: ___________________________________",
        ]
        for idx, line in enumerate(approval_lines):
            fig.text(0.18, 0.33 - (idx * 0.035), line, ha="left", va="center", fontsize=11, weight="bold" if idx == 0 else "normal")
        _pdf_page_base(fig, run_id, report_date, 1, total_pages)
        pdf.savefig(fig)
        plt.close(fig)

        # Page 2
        fig = plt.figure(figsize=(8.5, 11))
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        fig.text(0.08, 0.92, "What We Found", ha="left", va="top", fontsize=18, weight="bold", color="#1F4E79")
        p1 = (
            f"We compared {data['total_records']:,} employee records between your ADP system and Workday. "
            f"Of those, {data['approve_count']:,} records ({data['approve_pct']}%) matched cleanly and are ready to load. "
            f"{data['review_count']:,} records ({data['review_pct']}%) require human review before they can be loaded. "
            f"{data['reject_count']:,} records were flagged as possible wrong-person matches and have been blocked entirely."
        )
        p2 = (
            f"If you authorize this load, the following corrections will be applied to {data['approve_count']:,} Workday records:\n"
            f"- {data['job_org_count']:,} job title or department corrections\n"
            f"- {data['hire_date_count']:,} hire date corrections\n"
            f"- {data['status_count']:,} employment status corrections\n"
            f"- {data['salary_count']:,} salary corrections\n"
            f"No corrections will be applied to records in the review queue until a human reviewer clears them."
        )
        p3 = (
            f"The following records were automatically protected and will NOT receive any corrections regardless of authorization:\n"
            f"- {data['zero_salary_count']:,} active employees showing $0 salary in Workday (possible data entry error - requires investigation)\n"
            f"- {data['review_count']:,} records flagged for human review\n"
            f"- {data['reject_count']:,} records blocked as possible wrong-person matches"
        )
        _pdf_wrapped(fig, 0.08, 0.84, p1, width=88, fontsize=11)
        _pdf_wrapped(fig, 0.08, 0.64, p2, width=82, fontsize=11)
        _pdf_wrapped(fig, 0.08, 0.40, p3, width=82, fontsize=11)
        _pdf_page_base(fig, run_id, report_date, 2, total_pages)
        pdf.savefig(fig)
        plt.close(fig)

        # Page 3
        fig = plt.figure(figsize=(8.5, 11))
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        fig.text(0.08, 0.92, "Risks and Flags Identified", ha="left", va="top", fontsize=18, weight="bold", color="#1F4E79")
        risk_rows = _risk_rows(data)
        if risk_rows:
            table_ax = fig.add_axes([0.08, 0.58, 0.84, 0.26])
            table_ax.axis("off")
            table = table_ax.table(
                cellText=[[row[0], f"{row[1]:,}", row[2]] for row in risk_rows],
                colLabels=["Risk Type", "Count", "Action Required"],
                loc="upper left",
                cellLoc="left",
                colLoc="left",
                bbox=[0, 0, 1, 1],
            )
            table.auto_set_font_size(False)
            table.set_fontsize(10)
            for (r, c), cell in table.get_celld().items():
                if r == 0:
                    cell.set_text_props(weight="bold")
                    cell.set_facecolor("#F3F6F9")
        else:
            fig.text(0.08, 0.78, "No material risk rows were identified for this run.", ha="left", va="top", fontsize=11)
        gate_text = (
            f"The sanity gate result for this run was {data['gate_result']}. "
            f"If the gate result is FAIL, corrections have been blocked and this document "
            f"should not be signed until the gate failure reason has been investigated and resolved."
        )
        _pdf_wrapped(fig, 0.08, 0.48, gate_text, width=86, fontsize=11)
        if data["gate_result"] == "FAIL":
            ax.add_patch(Rectangle((0.08, 0.20), 0.84, 0.18, facecolor="#FDE9D9", edgecolor="#C0504D", linewidth=1.8))
            _pdf_wrapped(
                fig,
                0.10,
                0.35,
                "WARNING: SANITY GATE FAILED\nCorrections are currently blocked. Do not sign this document until the gate failure has been reviewed and resolved.\n"
                f"Gate failure reason: {'; '.join(data['gate_reasons']) if data['gate_reasons'] else 'No reason provided'}",
                width=82,
                fontsize=11,
                weight="bold",
            )
        _pdf_page_base(fig, run_id, report_date, 3, total_pages)
        pdf.savefig(fig)
        plt.close(fig)

        # Page 4
        fig = plt.figure(figsize=(8.5, 11))
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        fig.text(0.08, 0.93, "Authorization Record", ha="left", va="top", fontsize=18, weight="bold", color="#1F4E79")
        audit_trail = data["audit_trail"]
        input_files = audit_trail.get("input_files", {}) or {}
        old_file = input_files.get("old_system", {}) or {}
        new_file = input_files.get("new_system", {}) or {}
        auth_rows = [
            ["Run ID:", run_id],
            ["Run completed:", str(audit_trail.get("run_complete_timestamp") or "")],
            ["Engine version:", str(audit_trail.get("engine_version") or "")],
            ["Old system file:", f"{old_file.get('filename', '')} (SHA-256: {str(old_file.get('sha256') or '')[:16]}...)"],
            ["New system file:", f"{new_file.get('filename', '')} (SHA-256: {str(new_file.get('sha256') or '')[:16]}...)"],
            ["Total records processed:", f"{data['total_records']:,}"],
            ["Records approved:", f"{data['approve_count']:,}"],
            ["Records for review:", f"{data['review_count']:,}"],
            ["Records blocked:", f"{data['reject_count']:,}"],
            ["Corrections staged:", f"{data['corrections_staged_count']:,}"],
            ["Gate result:", data["gate_result"]],
        ]
        table_ax = fig.add_axes([0.08, 0.53, 0.84, 0.33])
        table_ax.axis("off")
        table = table_ax.table(cellText=auth_rows, colLabels=None, loc="upper left", cellLoc="left", bbox=[0, 0, 1, 1])
        table.auto_set_font_size(False)
        table.set_fontsize(9.5)
        for (r, c), cell in table.get_celld().items():
            if c == 0:
                cell.set_text_props(weight="bold")
                cell.set_facecolor("#F3F6F9")
        fig.text(0.08, 0.48, "What This Signature Authorizes", ha="left", va="top", fontsize=13, weight="bold", color="#1F4E79")
        _pdf_wrapped(
            fig,
            0.08,
            0.44,
            f"By signing this document, the Chief Human Resources Officer confirms they have reviewed the reconciliation "
            f"summary above and authorizes Data Whisperer to apply the staged corrections to {data['approve_count']:,} "
            f"Workday employee records. This authorization does not apply to records in the review queue or to records "
            f"that have been blocked. Those records require separate disposition before any corrections can be applied.",
            width=86,
            fontsize=11,
        )
        fig.text(0.08, 0.23, "Document Retention", ha="left", va="top", fontsize=13, weight="bold", color="#1F4E79")
        _pdf_wrapped(
            fig,
            0.08,
            0.19,
            "This signed document should be retained in accordance with your organization's HR records retention policy. "
            "The run ID above can be used to retrieve the full technical audit trail at any time.",
            width=86,
            fontsize=11,
        )
        _pdf_page_base(fig, run_id, report_date, 4, total_pages)
        pdf.savefig(fig)
        plt.close(fig)


def build_document(run_id: str, run_dir: Path, out_path: Path) -> None:
    data = _read_run_data(run_dir)
    _build_pdf(run_id, data, out_path)


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="Build CHRO approval PDF for a completed run.")
    parser.add_argument("--run-id", required=True, help="Run identifier")
    parser.add_argument("--run-dir", required=True, help="Run output directory")
    parser.add_argument("--out", required=True, help="Output PDF path")
    args = parser.parse_args(argv)

    run_dir = Path(args.run_dir)
    out_path = Path(args.out)
    build_document(args.run_id, run_dir, out_path)
    print(f"[build_chro_approval] saved: {out_path}")


if __name__ == "__main__":
    main()
