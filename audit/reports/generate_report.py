"""
generate_report.py - Audit Report Generator for the Data Whisperer reconciliation pipeline.

Reads data from audit/audit.db, audit/summary/*.csv, and audit/exports/out/wide_compare.csv
to produce a professional .docx audit report summarising the reconciliation run.

Report sections
---------------
  1. Executive Summary
  2. Match Quality Analysis
  3. Data Quality Findings
  4. Field Change Analysis  (Salary / Status / Hire Date / Job & Org)
  5. Priority Review Queue
  6. Rejected Matches (REJECT_MATCH tier)
  7. Recommendations

Dependencies
------------
  python-docx >= 1.1.0   (pip install python-docx)
  pandas
  PyYAML

Run:
    python3 audit/reports/generate_report.py [--db PATH] [--out PATH] [--wide PATH]
"""
from __future__ import annotations

import argparse
import json
import os
import re as _re
import sqlite3
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(
        "[error] python-docx not installed.\n"
        "  Run: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(2)

_HERE        = Path(__file__).resolve().parent
_SUMMARY_DIR = _HERE.parent / "summary"
sys.path.insert(0, str(_SUMMARY_DIR))

from gating import classify_all, salary_delta, _parse_confidence, _norm
from sanity_checks import detect_wave_dates

ROOT         = _HERE.parents[2]
DB_PATH      = ROOT / "audit" / "audit.db"
WIDE_CSV     = ROOT / "audit" / "exports" / "out" / "wide_compare.csv"
REVIEW_CSV   = ROOT / "audit" / "summary" / "review_queue.csv"
SANITY_GATE  = ROOT / "audit" / "summary" / "sanity_gate.json"
OUT_PATH     = _HERE / "audit_report.docx"

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
_BLUE_DARK    = RGBColor(0x2E, 0x75, 0xB6)   # section headings
_BLUE_LIGHT   = RGBColor(0xBD, 0xD7, 0xEE)   # table header fill
_RED          = RGBColor(0xC0, 0x00, 0x00)   # critical warnings
_ORANGE       = RGBColor(0xE2, 0x6B, 0x0A)   # caution
_GREEN        = RGBColor(0x37, 0x86, 0x45)   # positive
_GREY_LIGHT   = RGBColor(0xF2, 0xF2, 0xF2)   # alternating rows
_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Plain-English translation of internal engine flags
# ---------------------------------------------------------------------------

_REASON_MAP: dict[str, str] = {
    "hire_date:year_shift_with_other_mismatches":
        "Start date changed by a full year, plus other fields also changed",
    "hire_date:off_by_one_day_pattern":
        "Start date off by one day (likely a system convention difference - auto-approved)",
    "hire_date:year_shift_systematic":
        "Start date changed by one year - appears to be a systematic pattern (auto-approved)",
    "hire_date:off_by_one_year_systematic":
        "Start date off by one year across many records - auto-approved as systematic",
    "hire_date:off_by_one_year_pattern":
        "Start date off by exactly one year (likely a system import convention - auto-approved)",
    "worker_id_auto_approve":
        "Exact ID match - auto-approved",
    "pk_auto_approve":
        "Matched on name, date of birth, and last-4 identifier - auto-approved",
    "active_to_terminated":
        "Status changed from active to terminated - needs human review",
    "hire_date_wave":
        "Start date matches a bulk import date shared by many other employees - needs human review",
    "name_change_detected":
        "Last name differs between systems - needs human review to confirm same employee",
}

_MATCH_SOURCE_MAP: dict[str, str] = {
    "worker_id":      "Exact employee ID match",
    "pk":             "Matched on name, date of birth, and last-4",
    "last4_dob":      "Matched on last-4 SSN and date of birth",
    "dob_name":       "Matched on name and date of birth (high-risk tier)",
    "name_hire_date": "Matched on name and start date",
    "recon_id":       "Exact reconciliation ID match",
}


def _translate_reason(reason: str) -> str:
    """Translate an internal engine flag (or pipe-separated list) to plain English."""
    if not reason or str(reason).strip() in ("", "nan", "None"):
        return "No reason recorded"
    r = str(reason).strip()

    # Pipe-separated multi-reason (REVIEW records): translate each part
    if "|" in r:
        parts = [_translate_reason(p.strip()) for p in r.split("|") if p.strip()]
        return "; ".join(parts)

    # Strip "reject_match:" prefix (REJECT_MATCH tier reasons)
    r_inner = _re.sub(r"^reject_match:", "", r, flags=_re.IGNORECASE).strip()

    # Strip leading "fix_type:" prefix (REVIEW reasons like "salary:below_threshold")
    r_noprefix = _re.sub(r"^[a-z_]+:", "", r_inner).strip()

    # dob_name_low_confidence (0.600<0.75) - extract actual score AND minimum threshold
    m = _re.match(r"dob_name_low_confidence\s*\(?([\d.]+)<([\d.]+)", r_inner)
    if m:
        try:
            pct_score  = int(round(float(m.group(1)) * 100))
            pct_thresh = int(round(float(m.group(2)) * 100))
        except ValueError:
            pct_score, pct_thresh = 0, 75
        return (
            f"Matched on name and date of birth only - confidence score of {pct_score}% "
            f"fell below the {pct_thresh}% minimum required to trust the match."
        )

    # Legacy: dob_name_low_confidence (0.600) without threshold value
    m = _re.match(r"dob_name_low_confidence\s*\(?([\d.]+)", r_inner)
    if m:
        try:
            pct = int(round(float(m.group(1)) * 100))
        except ValueError:
            pct = 0
        return f"Matched on name and date of birth only - confidence too low to trust ({pct}%)"

    # fuzzy_extreme_salary_ratio (2.5000>2.5) - wrong-person signal from salary mismatch
    m = _re.match(r"fuzzy_extreme_salary_ratio\s*\(?([\d.]+)", r_inner)
    if m:
        try:
            ratio = float(m.group(1))
        except ValueError:
            ratio = 2.5
        return (
            f"Salary is {ratio:.1f}x different between systems on a fuzzy match - "
            f"flagged as possible wrong-person pairing"
        )

    # salary_ratio_extreme (2.0000 outside [0.85, 1.15]) - large salary change on known worker
    m = (_re.match(r"salary_ratio_extreme\s*\(?([\d.]+)", r_noprefix) or
         _re.match(r"salary_ratio_extreme\s*\(?([\d.]+)", r))
    if m:
        try:
            ratio = float(m.group(1))
        except ValueError:
            ratio = 2.0
        if ratio > 1.5:
            return "Salary more than doubled between systems - needs human review"
        if ratio < 0.5:
            return "Salary dropped by more than half between systems - needs human review"
        return "Salary changed significantly between systems - needs human review"

    # below_threshold (0.82<0.97) - confidence below the required minimum for this field type
    m = _re.match(r"below_threshold\s*\(?([\d.]+)<([\d.]+)", r_noprefix)
    if m:
        try:
            pct_score  = int(round(float(m.group(1)) * 100))
            pct_thresh = int(round(float(m.group(2)) * 100))
        except ValueError:
            pct_score, pct_thresh = 0, 95
        return (
            f"Confidence score of {pct_score}% fell below the {pct_thresh}% "
            f"minimum required for this field - needs human review"
        )

    # low_confidence (0.82) - single value fallback
    m = _re.match(r"low_confidence\s*\(?([\d.]+)", r_noprefix)
    if m:
        try:
            pct = int(round(float(m.group(1)) * 100))
        except ValueError:
            pct = 0
        return f"Confidence score too low ({pct}%) - needs human review"

    # active_to_terminated (active->terminated)
    if _re.match(r"active_to_terminated", r_noprefix):
        return "Status changed from active to terminated - needs human review"

    # name_change_detected (old_last -> new_last) - parametric form
    m = _re.match(r"name_change_detected\s*\(([^)]+)\)", r_noprefix)
    if m:
        parts = m.group(1).strip()
        return f"Last name changed ({parts}) - needs human review to confirm same employee"

    # Exact map lookups: try r_inner, r_noprefix, then full r
    for candidate in (r_inner, r_noprefix, r):
        if candidate in _REASON_MAP:
            return _REASON_MAP[candidate]

    # Prefix map lookup
    for key, val in _REASON_MAP.items():
        if r_inner.startswith(key) or r_noprefix.startswith(key) or r.startswith(key):
            return val

    # Generic fallback: humanise snake_case from the most-stripped form
    display = r_noprefix or r_inner or r
    return (display.replace("hire_date:", "start date: ")
                   .replace("_", " ")
                   .replace("  ", " ")
                   .strip()
                   .capitalize())


def _translate_match_source(source: str) -> str:
    """Translate a match_source key to a plain-English label."""
    return _MATCH_SOURCE_MAP.get(str(source).strip().lower(), str(source))


def _display_name(row: dict) -> str:
    """Return 'First Last' display name from name components or full_name_norm fallback.

    Priority:
      1. first_name_norm + last_name_norm  (title-cased, clean)
      2. old_full_name_norm                (title-cased)
      3. pair_id                           (fallback)
    """
    first = str(row.get("old_first_name_norm") or "").strip()
    last  = str(row.get("old_last_name_norm")  or "").strip()
    if first and last:
        # Title-case each component (handles hyphenated names gracefully)
        parts = [w.capitalize() for w in first.split()]
        lparts = [w.capitalize() for w in last.split()]
        return " ".join(parts) + " " + " ".join(lparts)
    # Fallback: use full_name_norm and title-case it
    full = str(row.get("old_full_name_norm") or "").strip()
    if full:
        return " ".join(w.capitalize() for w in full.split())
    return str(row.get("pair_id", "-"))


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set cell background colour (XML shading element)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = _BLUE_DARK
            run.font.bold = True


def _kv_table(doc: Document, rows: list[tuple]) -> None:
    """Write a 2-column key-value table with alternating row colours."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        cells = tbl.rows[i].cells
        cells[0].text = str(key)
        cells[1].text = str(val)
        # Bold key column
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
        # Alternating background
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(cells[0], bg)
        _set_cell_bg(cells[1], bg)
    # Column widths
    tbl.columns[0].width = Inches(3.0)
    tbl.columns[1].width = Inches(2.5)


def _data_table(doc: Document, headers: list[str], rows_data: list[list]) -> None:
    """Write a header + data table."""
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = _WHITE
        _set_cell_bg(hdr_cells[i], "2E75B6")

    # Data rows
    for r_i, row_vals in enumerate(rows_data):
        cells = tbl.rows[r_i + 1].cells
        for c_i, val in enumerate(row_vals):
            cells[c_i].text = str(val) if val is not None else ""
            bg = "F2F2F2" if r_i % 2 == 0 else "FFFFFF"
            _set_cell_bg(cells[c_i], bg)


def _callout(doc: Document, text: str, level: str = "warning") -> None:
    """Add a callout paragraph (bold, coloured text)."""
    p    = doc.add_paragraph()
    run  = p.add_run(f"{'⚠' if level == 'warning' else '✓'}  {text}")
    run.font.bold = True
    run.font.color.rgb = _RED if level == "critical" else (_ORANGE if level == "warning" else _GREEN)


# ---------------------------------------------------------------------------
# Data loader
# ---------------------------------------------------------------------------

def _load_data(db_path: Path, wide_path: Path) -> pd.DataFrame:
    """Load matched pairs with gating results from wide_compare CSV or DB."""
    if wide_path.exists():
        df = pd.read_csv(str(wide_path))
        print(f"[generate_report] loaded {len(df):,} rows from {wide_path.name}")
        return df

    print(f"[generate_report] {wide_path.name} not found - computing from DB ...")
    con = sqlite3.connect(str(db_path))
    try:
        mp = pd.read_sql_query("SELECT * FROM matched_pairs", con)
    finally:
        con.close()

    if "confidence" not in mp.columns:
        mp["confidence"] = None

    all_rows   = mp.to_dict(orient="records")
    wave_dates = detect_wave_dates(all_rows)

    out_rows = []
    for r in all_rows:
        result    = classify_all(r, wave_dates=wave_dates)
        fix_types = result["fix_types"]
        sal_d     = salary_delta(r)
        out_rows.append({
            "pair_id":           r.get("pair_id", ""),
            "match_source":      r.get("match_source", ""),
            "confidence":        r.get("confidence"),
            "action":            result["action"],
            "reason":            result["reason"],
            "fix_types":         "|".join(fix_types),
            "old_full_name_norm":r.get("old_full_name_norm", ""),
            "new_full_name_norm":r.get("new_full_name_norm", ""),
            "old_salary":        r.get("old_salary"),
            "new_salary":        r.get("new_salary"),
            "old_worker_status": r.get("old_worker_status", ""),
            "new_worker_status": r.get("new_worker_status", ""),
            "old_hire_date":     r.get("old_hire_date", ""),
            "new_hire_date":     r.get("new_hire_date", ""),
            "old_position":      r.get("old_position", ""),
            "new_position":      r.get("new_position", ""),
            "salary_delta":      sal_d,
        })
    print(f"[generate_report] computed gating for {len(out_rows):,} rows from DB")
    return pd.DataFrame(out_rows)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _section_executive_summary(doc: Document, df: pd.DataFrame) -> None:
    _add_heading(doc, "1. Executive Summary", 1)

    total     = len(df)
    n_approve = int((df["action"] == "APPROVE").sum())     if "action" in df.columns else 0
    n_review  = int((df["action"] == "REVIEW").sum())      if "action" in df.columns else 0
    n_reject  = int((df["action"] == "REJECT_MATCH").sum()) if "action" in df.columns else 0

    # Active/$0 count
    active_zero = 0
    if "new_worker_status" in df.columns and "new_salary" in df.columns:
        _am = df["new_worker_status"].fillna("").str.strip().str.lower() == "active"
        _sn = pd.to_numeric(
            df["new_salary"].astype(str).str.replace(",", "").str.replace("$", ""),
            errors="coerce",
        )
        active_zero = int((_am & (_sn.isna() | (_sn == 0))).sum())

    # -----------------------------------------------------------------------
    # Plain-language opening - readable by non-HR managers in 60 seconds
    # -----------------------------------------------------------------------
    doc.add_paragraph(
        f"We compared {total:,} employee records between the source system and the new "
        f"system. Here is what we found:"
    )

    bullets = [
        f"{n_approve:,} records look correct and are ready to load - no action needed.",
        f"{n_review:,} records need a human to review them before they go in.",
    ]
    if n_reject > 0:
        bullets.append(
            f"{n_reject:,} records were flagged as possible wrong-person matches "
            f"and blocked entirely."
        )
    if active_zero > 0:
        bullets.append(
            f"{active_zero:,} employees are showing $0 salary in the new system - "
            f"this needs to be fixed before anyone gets paid."
        )

    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # What happens next
    # -----------------------------------------------------------------------
    _add_heading(doc, "What happens next", 2)
    doc.add_paragraph(
        f"The correction files attached to this report are ready to load into the new system "
        f"for all {n_approve:,} auto-approved records. Before loading, a reviewer must work "
        f"through the {n_review:,} records in the review queue and confirm each one. "
        f"Once the review queue is cleared, a final corrections run can be executed."
    )

    if active_zero > 0:
        _callout(doc,
            f"Action required before payroll: {active_zero:,} active employees have $0 salary "
            f"in the new system. Salary corrections for these records are blocked until the "
            f"source data is corrected.",
            level="critical",
        )

    if n_reject > 0:
        _callout(doc,
            f"{n_reject:,} employee records were blocked from corrections entirely - these appear "
            f"to be wrong-person matches and need manual investigation. Do not load corrections "
            f"for these records.",
            level="warning",
        )


def _section_match_quality(doc: Document, df: pd.DataFrame) -> None:
    _add_heading(doc, "2. Match Quality Analysis", 1)

    if "match_source" not in df.columns:
        doc.add_paragraph("(match_source column not available)")
        return

    total = len(df)
    _add_heading(doc, "2.1 Match Source Breakdown", 2)
    src_counts = df["match_source"].value_counts()
    table_rows = []
    for src, cnt in src_counts.items():
        pct     = f"{cnt/total*100:.1f}%"
        src_str = str(src) if src else "(blank)"
        table_rows.append([src_str, f"{cnt:,}", pct])
    _data_table(doc,
        ["Match Source", "Count", "% of Total"],
        table_rows
    )
    doc.add_paragraph()

    if "confidence" in df.columns:
        _add_heading(doc, "2.2 Confidence Distribution", 2)
        conf_num    = pd.to_numeric(df["confidence"], errors="coerce")
        n_exact     = int((conf_num == 1.0).sum())
        n_high      = int(((conf_num >= 0.97) & (conf_num < 1.0)).sum())
        n_medium    = int(((conf_num >= 0.80) & (conf_num < 0.97)).sum())
        n_low       = int(((conf_num < 0.80) & conf_num.notna()).sum())
        n_missing   = int(conf_num.isna().sum())
        _data_table(doc,
            ["Confidence Band", "Count", "% of Total"],
            [
                ["Exact (1.00)",          f"{n_exact:,}",   f"{n_exact/total*100:.1f}%"],
                ["High (0.97-0.99)",      f"{n_high:,}",    f"{n_high/total*100:.1f}%"],
                ["Medium (0.80-0.96)",    f"{n_medium:,}",  f"{n_medium/total*100:.1f}%"],
                ["Low (< 0.80)",          f"{n_low:,}",     f"{n_low/total*100:.1f}%"],
                ["Missing / not scored",  f"{n_missing:,}", f"{n_missing/total*100:.1f}%"],
            ]
        )
        doc.add_paragraph()

        if n_low > 0:
            _callout(doc,
                f"{n_low:,} pairs have low confidence scores (< 0.80). "
                f"These require careful human review before any corrections are applied.",
                level="warning"
            )


def _section_data_quality(doc: Document, df: pd.DataFrame) -> None:
    _add_heading(doc, "3. Data Quality Findings", 1)

    # Active/$0 salary
    _add_heading(doc, "3.1 Active Employees with $0 Salary", 2)
    if "new_worker_status" in df.columns and "new_salary" in df.columns:
        active_mask = df["new_worker_status"].fillna("").str.strip().str.lower() == "active"
        new_sal_num = pd.to_numeric(
            df["new_salary"].astype(str).str.replace(",","").str.replace("$",""),
            errors="coerce"
        )
        az_df = df[active_mask & (new_sal_num.isna() | (new_sal_num == 0))].copy()
        if az_df.empty:
            p = doc.add_paragraph("No active employees with $0 or missing salary detected. ")
            run = p.runs[0]
            run.font.color.rgb = _GREEN
        else:
            _callout(doc,
                f"{len(az_df):,} active employees have $0 or missing salary. "
                "Salary corrections for these records have been blocked from the corrections pipeline.",
                level="critical"
            )
            sample_cols = [c for c in ["pair_id","old_full_name_norm","old_salary","new_salary"] if c in az_df.columns]
            if sample_cols:
                sample = az_df[sample_cols].head(10)
                _data_table(doc, sample_cols, sample.values.tolist())
    else:
        doc.add_paragraph("(Status/salary columns not available for this check)")
    doc.add_paragraph()

    # Wave dates
    _add_heading(doc, "3.2 Hire Date Wave Detection", 2)
    if "new_hire_date" in df.columns and "reason" in df.columns:
        wave_rows = df[df["reason"].fillna("").str.contains("hire_date_wave", na=False)]
        if wave_rows.empty:
            p = doc.add_paragraph("No hire date wave patterns detected. ")
            p.runs[0].font.color.rgb = _GREEN
        else:
            wave_date_vals = df["new_hire_date"].value_counts()
            wave_dates_top = wave_date_vals[wave_date_vals > len(df) * 0.01].head(10)
            _callout(doc,
                f"{len(wave_rows):,} records share hire dates that appear in ≥ 1% of all pairs "
                f"- indicative of a bulk import. These have been routed to REVIEW.",
                level="warning"
            )
            if not wave_dates_top.empty:
                _data_table(doc,
                    ["Hire Date", "Count", "% of Total"],
                    [[d, f"{c:,}", f"{c/len(df)*100:.1f}%"] for d, c in wave_dates_top.items()]
                )
    else:
        doc.add_paragraph("(Hire date / reason columns not available)")
    doc.add_paragraph()

    # Blocked records (REJECT_MATCH)
    _add_heading(doc, "3.3 Blocked Records - Possible Wrong-Person Matches", 2)
    if "action" in df.columns:
        rej_df = df[df["action"] == "REJECT_MATCH"]
        if rej_df.empty:
            p = doc.add_paragraph("No blocked records detected. All pairings appear valid. ")
            p.runs[0].font.color.rgb = _GREEN
        else:
            _callout(doc,
                f"{len(rej_df):,} pairs were blocked - the matching engine detected these are "
                f"likely wrong-person pairings. No corrections will be applied to these records.",
                level="critical",
            )
            if "reason" in rej_df.columns:
                rej_reasons = rej_df["reason"].value_counts().head(5)
                _data_table(doc,
                    ["Why They Were Blocked", "Count"],
                    [[_translate_reason(str(r)), f"{c:,}"] for r, c in rej_reasons.items()]
                )
    doc.add_paragraph()


def _section_field_changes(doc: Document, df: pd.DataFrame) -> None:
    _add_heading(doc, "4. Field Change Analysis", 1)

    if "fix_types" not in df.columns:
        doc.add_paragraph("(fix_types column not available)")
        return

    total = len(df)

    # Summary table
    change_rows = []
    for ft, label in [
        ("salary",    "Salary"),
        ("payrate",   "Payrate"),
        ("status",    "Worker Status"),
        ("hire_date", "Hire Date"),
        ("job_org",   "Job / Organization"),
    ]:
        cnt = int(df["fix_types"].str.contains(ft, na=False).sum())
        change_rows.append([label, f"{cnt:,}", f"{cnt/total*100:.1f}%"])
    _data_table(doc, ["Change Type", "Count", "% of Total"], change_rows)
    doc.add_paragraph()

    # Salary details
    _add_heading(doc, "4.1 Salary Changes", 2)
    if "salary_delta" in df.columns:
        sal_df_all = df[df["fix_types"].str.contains("salary", na=False)].copy()

        # Exclude Active/$0 records: these are data quality issues (missing salary
        # in source data), not real salary changes.  Including them collapses
        # mean/median dramatically - e.g. $0 for a $50k worker gives delta -$50k.
        # Same exclusion used by build_workbook.py validate_active_zero_salary().
        n_excluded = 0
        if "new_worker_status" in sal_df_all.columns and "new_salary" in sal_df_all.columns:
            _am = sal_df_all["new_worker_status"].fillna("").str.strip().str.lower() == "active"
            _sn = pd.to_numeric(
                sal_df_all["new_salary"].astype(str).str.replace(",", "").str.replace("$", ""),
                errors="coerce",
            )
            _az_mask   = _am & (_sn.isna() | (_sn == 0))
            n_excluded = int(_az_mask.sum())
            sal_df     = sal_df_all[~_az_mask].copy()
        else:
            sal_df = sal_df_all.copy()

        sal_delta = pd.to_numeric(sal_df["salary_delta"], errors="coerce").dropna()
        sal_delta = sal_delta[sal_delta != 0]   # only rows with an actual change

        if len(sal_delta) > 0:
            n_increase = int((sal_delta > 0).sum())
            n_decrease = int((sal_delta < 0).sum())
            kv_rows = [("Salary changes (total)",  f"{len(sal_df_all):,}")]
            if n_excluded > 0:
                kv_rows.append((
                    "  - Active/$0 excl. from stats",
                    f"{n_excluded:,}  (data quality - not real changes)",
                ))
            kv_rows += [
                ("  - Included in stats",     f"{len(sal_df):,}"),
                ("Increases",                 f"{n_increase:,}"),
                ("Decreases",                 f"{n_decrease:,}"),
                ("Mean delta",                f"${sal_delta.mean():,.2f}"),
                ("Median delta",              f"${sal_delta.median():,.2f}"),
                ("Largest increase",          f"${sal_delta.max():,.2f}"),
                ("Largest decrease",          f"${sal_delta.min():,.2f}"),
            ]
            _kv_table(doc, kv_rows)
        else:
            doc.add_paragraph("No parseable salary deltas found.")
    else:
        doc.add_paragraph("(salary_delta column not available)")
    doc.add_paragraph()

    # Status changes
    _add_heading(doc, "4.2 Status Changes", 2)
    if "old_worker_status" in df.columns and "new_worker_status" in df.columns:
        st_df = df[df["fix_types"].str.contains("status", na=False)].copy()
        if st_df.empty:
            doc.add_paragraph("No status changes detected.")
        else:
            transitions = (
                st_df["old_worker_status"].fillna("blank").str.lower().str.strip()
                + " → "
                + st_df["new_worker_status"].fillna("blank").str.lower().str.strip()
            ).value_counts().head(10)
            _data_table(doc,
                ["Transition", "Count"],
                [[t, f"{c:,}"] for t, c in transitions.items()]
            )
    doc.add_paragraph()

    # Hire date changes - plain-English summary (no technical pattern table)
    _add_heading(doc, "4.3 Hire Date Changes", 2)
    hd_df = df[df["fix_types"].str.contains("hire_date", na=False)].copy() \
        if "fix_types" in df.columns else pd.DataFrame()
    if hd_df.empty:
        doc.add_paragraph("No hire date differences detected.")
    else:
        n_total = len(hd_df)

        # Categorise by pattern
        n_auto_approved = 0
        n_year_review   = 0
        n_systematic    = 0
        reason_col = "hire_date_pattern" if "hire_date_pattern" in hd_df.columns else \
                     "reason"            if "reason"            in hd_df.columns else None
        if reason_col:
            vals = hd_df[reason_col].fillna("")
            n_auto_approved = int(vals.str.contains("off_by_one_day", na=False).sum())
            n_year_review   = int(vals.str.contains("year_shift_with_other", na=False).sum())
            n_systematic    = int(
                vals.str.contains("year_shift_systematic|off_by_one_year_systematic", na=False).sum()
            )

        n_other = max(0, n_total - n_auto_approved - n_year_review - n_systematic)

        parts: list[str] = []
        if n_auto_approved > 0:
            parts.append(
                f"{n_auto_approved:,} were automatically approved because they matched known "
                f"system conversion patterns (such as off-by-one-day differences)"
            )
        if n_year_review > 0:
            parts.append(
                f"{n_year_review:,} had a full year shift combined with other changes "
                f"and were sent to review"
            )
        if n_systematic > 0:
            parts.append(
                f"{n_systematic:,} appear to be a systematic date format difference "
                f"and were auto-approved"
            )
        if n_other > 0:
            parts.append(f"{n_other:,} had other hire date differences flagged for review")

        if parts:
            doc.add_paragraph(f"Of {n_total:,} hire date differences found: " + "; ".join(parts) + ".")
        else:
            doc.add_paragraph(f"A total of {n_total:,} records had hire date differences.")
    doc.add_paragraph()


def _section_review_queue(doc: Document, df: pd.DataFrame) -> None:
    _add_heading(doc, "5. Priority Review Queue", 1)

    if "action" not in df.columns:
        doc.add_paragraph("(action column not available)")
        return

    review_df    = df[df["action"] == "REVIEW"].copy()
    total_review = len(review_df)

    doc.add_paragraph(
        f"A total of {total_review:,} records need a human reviewer to look at them before "
        f"corrections can be applied. The table below shows the 10 highest-priority items."
    )

    if total_review == 0:
        _callout(doc,
            "No records require human review. All approved records are ready to load.",
            level="ok",
        )
        return

    # Sort by priority if available, otherwise take first 10
    if "priority_score" in review_df.columns:
        top_review = review_df.sort_values("priority_score", ascending=False).head(10)
    else:
        top_review = review_df.head(10)

    # Build plain-English table: Employee Name | Why It Needs Review | Fields Changed
    has_reason = "reason"    in top_review.columns
    has_fix    = "fix_types" in top_review.columns
    has_name   = "old_full_name_norm" in top_review.columns

    rows_data = []
    for _, row in top_review.iterrows():
        name     = _display_name(row.to_dict())  # "First Last" format
        reason   = _translate_reason(str(row.get("reason", ""))) if has_reason else "-"
        fix_typs = str(row.get("fix_types", "")).replace("|", ", ") if has_fix else "-"
        rows_data.append([name, reason, fix_typs])

    _data_table(doc, ["Employee", "Why It Needs Review", "Fields Changed"], rows_data)

    if total_review > 10:
        doc.add_paragraph(
            f"Showing 10 of {total_review:,} total review items. "
            f"See the review_queue.csv file for the complete list."
        )

    # Summary by change type
    if has_fix and len(review_df) > 0:
        doc.add_paragraph()
        _add_heading(doc, "Review Queue by Change Type", 2)
        ft_counts: dict[str, int] = {}
        for _, row in review_df.iterrows():
            for ft in str(row.get("fix_types", "")).split("|"):
                ft = ft.strip()
                if ft:
                    ft_counts[ft] = ft_counts.get(ft, 0) + 1
        _data_table(doc,
            ["Change Type", "Records Needing Review"],
            [[ft, f"{c:,}"] for ft, c in sorted(ft_counts.items(), key=lambda x: -x[1])],
        )
    doc.add_paragraph()


def _section_rejected_matches(doc: Document, df: pd.DataFrame) -> None:
    _add_heading(doc, "6. Blocked Records - Possible Wrong-Person Matches", 1)

    if "action" not in df.columns:
        doc.add_paragraph("(action column not available)")
        return

    rej_df = df[df["action"] == "REJECT_MATCH"].copy()

    if rej_df.empty:
        _callout(doc,
            "No records were blocked. All pairings appear to be correct person matches.",
            level="ok",
        )
        return

    doc.add_paragraph(
        f"{len(rej_df):,} records were blocked from corrections because the matching engine "
        f"detected they were likely wrong-person pairings - the system matched an employee "
        f"to someone else in the other file. These records have been completely excluded "
        f"from all correction files and must be investigated and re-matched manually."
    )
    doc.add_paragraph(
        "The most common reason for a block is a name-and-date-of-birth-only match where the "
        "confidence score fell below the minimum acceptable threshold. These low-confidence "
        "matches carry a real risk of applying salary, status, or hire-date changes to the "
        "wrong person."
    )

    if "match_source" in rej_df.columns:
        src_counts = rej_df["match_source"].value_counts()
        _data_table(doc,
            ["How They Were Originally Matched", "Count Blocked"],
            [[_translate_match_source(str(s)), f"{c:,}"] for s, c in src_counts.items()],
        )
    doc.add_paragraph()


def _section_recommendations(doc: Document, df: pd.DataFrame) -> None:
    _add_heading(doc, "7. Recommendations", 1)

    total        = len(df)
    n_rev        = int((df["action"] == "REVIEW").sum())      if "action" in df.columns else 0
    n_rej        = int((df["action"] == "REJECT_MATCH").sum()) if "action" in df.columns else 0
    approve_rate = (total - n_rev - n_rej) / total if total > 0 else 0.0

    active_zero_count = 0
    if "new_worker_status" in df.columns and "new_salary" in df.columns:
        am = df["new_worker_status"].fillna("").str.strip().str.lower() == "active"
        sn = pd.to_numeric(
            df["new_salary"].astype(str).str.replace(",", "").str.replace("$", ""),
            errors="coerce",
        )
        active_zero_count = int((am & (sn.isna() | (sn == 0))).sum())

    rec_items: list[str] = []

    if active_zero_count > 0:
        rec_items.append(
            f"Fix {active_zero_count:,} active employees who show $0 salary in the new system. "
            f"This is a data extraction issue - these employees had salary data in the old system "
            f"that did not transfer correctly. Salary corrections for these records are blocked "
            f"until the underlying data problem is resolved."
        )

    if n_rej > 0:
        rec_items.append(
            f"Investigate {n_rej:,} blocked records. These appear to be wrong-person matches "
            f"where the engine paired an employee with the wrong record in the other file. "
            f"Each one needs to be manually reviewed and either re-matched or excluded from "
            f"the migration entirely."
        )

    if n_rev > 0:
        rec_items.append(
            f"Work through the review queue before running corrections. {n_rev:,} records need "
            f"a reviewer to look at them and confirm they are correct. Some corrections were "
            f"held and not applied automatically - these are in the held_corrections file and "
            f"require manual review and approval before they can be loaded into the new system."
        )

    if approve_rate < 0.80:
        rec_items.append(
            f"The auto-approval rate is {approve_rate:.1%}, which is below the recommended "
            f"80% target. This means more records than expected need human review. Check whether "
            f"the source data extract is complete and whether any fields are missing or malformed."
        )

    if not rec_items:
        rec_items.append(
            "All data quality checks passed. The correction files are ready to load into the "
            "new system. Review and apply the corrections in the output folder to complete "
            "the migration."
        )

    for item in rec_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Node.js generator support
# ---------------------------------------------------------------------------

def _active_zero_mask(df: pd.DataFrame) -> "pd.Series":
    """Return boolean mask: Active workers where new_salary is 0 or null."""
    if "new_worker_status" not in df.columns or "new_salary" not in df.columns:
        return pd.Series([False] * len(df), index=df.index)
    am = df["new_worker_status"].fillna("").str.strip().str.lower() == "active"
    sn = pd.to_numeric(
        df["new_salary"].astype(str).str.replace(",", "").str.replace("$", ""),
        errors="coerce",
    )
    return am & (sn.isna() | (sn == 0))


def _serialize_run_data(df: pd.DataFrame, db_path: Path, wide_path: Path) -> dict:
    """Serialize the full run dataset into a JSON-serialisable dict for build_report.js."""
    total  = len(df)
    az_mask = _active_zero_mask(df)

    # Actions
    def _n(col_val: str) -> int:
        return int((df["action"] == col_val).sum()) if "action" in df.columns else 0
    actions = {"APPROVE": _n("APPROVE"), "REVIEW": _n("REVIEW"), "REJECT_MATCH": _n("REJECT_MATCH")}

    # Match sources
    match_sources = (
        [[str(s), int(c)] for s, c in df["match_source"].value_counts().items()]
        if "match_source" in df.columns else []
    )

    # Confidence bands
    confidence_bands: dict = {}
    if "confidence" in df.columns:
        cn = pd.to_numeric(df["confidence"], errors="coerce")
        confidence_bands = {
            "exact":   int((cn == 1.0).sum()),
            "high":    int(((cn >= 0.97) & (cn < 1.0)).sum()),
            "medium":  int(((cn >= 0.80) & (cn < 0.97)).sum()),
            "low":     int(((cn < 0.80) & cn.notna()).sum()),
            "missing": int(cn.isna().sum()),
        }

    # Active/$0 sample (up to 10 rows for table in report)
    active_zero_count = int(az_mask.sum())
    active_zero_sample: list = []
    if active_zero_count > 0:
        az_df = df[az_mask]
        sample_cols = [c for c in ["pair_id", "old_full_name_norm", "old_salary", "new_salary"] if c in az_df.columns]
        for _, row in az_df[sample_cols].head(10).iterrows():
            active_zero_sample.append([str(row.get(c, "")) for c in sample_cols])

    # Salary stats (Active/$0 excluded)
    salary: dict | None = None
    if "salary_delta" in df.columns and "fix_types" in df.columns:
        sal_all  = df[df["fix_types"].str.contains("salary", na=False)].copy()
        az_in    = az_mask.reindex(sal_all.index, fill_value=False)
        sal_incl = sal_all[~az_in]
        sal_d    = pd.to_numeric(sal_incl["salary_delta"], errors="coerce").dropna()
        sal_d    = sal_d[sal_d != 0]
        salary = {
            "total":                   int(len(sal_all)),
            "n_excluded_active_zero":  int(az_in.sum()),
            "n_included":              int(len(sal_incl)),
            "n_increase":              int((sal_d > 0).sum())         if len(sal_d) > 0 else 0,
            "n_decrease":              int((sal_d < 0).sum())         if len(sal_d) > 0 else 0,
            "mean_delta":              round(float(sal_d.mean()), 2)   if len(sal_d) > 0 else 0.0,
            "median_delta":            round(float(sal_d.median()), 2) if len(sal_d) > 0 else 0.0,
            "max_increase":            round(float(sal_d.max()), 2)    if len(sal_d) > 0 else 0.0,
            "max_decrease":            round(float(sal_d.min()), 2)    if len(sal_d) > 0 else 0.0,
        }

    # Status transitions
    status_transitions: list = []
    if all(c in df.columns for c in ["old_worker_status", "new_worker_status", "fix_types"]):
        st_df = df[df["fix_types"].str.contains("status", na=False)]
        trans = (
            st_df["old_worker_status"].fillna("blank").str.lower().str.strip()
            + " → "
            + st_df["new_worker_status"].fillna("blank").str.lower().str.strip()
        ).value_counts().head(10)
        status_transitions = [[str(t), int(c)] for t, c in trans.items()]

    # Hire date stats
    hire_date_stats: dict = {}
    if "fix_types" in df.columns:
        hd_df = df[df["fix_types"].str.contains("hire_date", na=False)]
        hire_date_stats["total"] = int(len(hd_df))
        if "reason" in df.columns:
            hire_date_stats["n_wave"] = int(df["reason"].fillna("").str.contains("hire_date_wave").sum())
        if "hire_date_pattern" in hd_df.columns:
            pats = hd_df["hire_date_pattern"].fillna("").replace("", "none").value_counts().head(8)
            hire_date_stats["patterns"] = [[str(p), int(c)] for p, c in pats.items()]

    # Wave dates
    wave_dates: list = []
    if "new_hire_date" in df.columns and "reason" in df.columns:
        wr = df[df["reason"].fillna("").str.contains("hire_date_wave", na=False)]
        if not wr.empty:
            wave_dates = [[str(d), int(c)] for d, c in wr["new_hire_date"].value_counts().head(10).items()]

    # Reject matches
    reject_matches: dict = {"total": 0, "reasons": [], "by_source": []}
    if "action" in df.columns:
        rej_df = df[df["action"] == "REJECT_MATCH"]
        reject_matches["total"] = int(len(rej_df))
        if not rej_df.empty:
            if "reason" in rej_df.columns:
                reject_matches["reasons"] = [[str(r), int(c)] for r, c in rej_df["reason"].value_counts().head(5).items()]
            if "match_source" in rej_df.columns:
                reject_matches["by_source"] = [[str(s), int(c)] for s, c in rej_df["match_source"].value_counts().items()]

    # Review queue
    review_queue: dict = {"total": 0, "top_items": [], "by_fix_type": []}
    if "action" in df.columns:
        rev_df = df[df["action"] == "REVIEW"].copy()
        review_queue["total"] = int(len(rev_df))
        if "priority_score" in rev_df.columns:
            top = rev_df.sort_values("priority_score", ascending=False).head(20)
        else:
            top = rev_df.head(20)
        if len(top) > 0:
            # top_items format: [display_name, reason, fix_types]
            # display_name uses first/last components when available, falls back to full_name_norm
            items = []
            for _, row in top.iterrows():
                row_dict = row.to_dict()
                items.append([
                    _display_name(row_dict),
                    str(row_dict.get("reason", "")),
                    str(row_dict.get("fix_types", "")),
                ])
            review_queue["top_items"] = items
        if "fix_types" in rev_df.columns:
            ft_c: dict = {}
            for _, row in rev_df.iterrows():
                for ft in str(row.get("fix_types", "")).split("|"):
                    if ft:
                        ft_c[ft] = ft_c.get(ft, 0) + 1
            review_queue["by_fix_type"] = [[ft, c] for ft, c in sorted(ft_c.items(), key=lambda x: -x[1])]

    # Computed findings list
    sev_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "BUG": 3, "INFO": 4}
    findings: list = []
    if active_zero_count > 0:
        findings.append({"severity": "CRITICAL", "count": active_zero_count,
                         "title":  "Active Employees with $0 Salary",
                         "impact": "Salary corrections blocked; source data correction required"})
    if reject_matches["total"] > 0:
        findings.append({"severity": "HIGH", "count": reject_matches["total"],
                         "title":  "Wrong-Person Pairings (REJECT_MATCH)",
                         "impact": "Excluded from all corrections; require manual investigation"})
    if actions["REVIEW"] > 0:
        findings.append({"severity": "HIGH", "count": actions["REVIEW"],
                         "title":  "Pairs Requiring Human Review",
                         "impact": "Corrections held until reviewer approves"})
    if "confidence" in df.columns:
        cn     = pd.to_numeric(df["confidence"], errors="coerce")
        n_low  = int(((cn < 0.70) & cn.notna()).sum())
        if n_low > 0:
            findings.append({"severity": "HIGH", "count": n_low,
                             "title":  "Low-Confidence Fuzzy Matches (< 0.70)",
                             "impact": "High risk of wrong-person corrections; verify before applying"})
    findings.sort(key=lambda f: sev_order.get(f["severity"], 99))

    # Change summary
    change_summary: list = []
    if "fix_types" in df.columns:
        for ft, label in [("salary", "Salary"), ("payrate", "Payrate"),
                          ("status", "Worker Status"), ("hire_date", "Hire Date"),
                          ("job_org", "Job / Organization")]:
            cnt = int(df["fix_types"].str.contains(ft, na=False).sum())
            change_summary.append([label, f"{cnt:,}", f"{cnt/total*100:.1f}%" if total > 0 else "0.0%"])

    # Sanity gate - try per-run path first, then promotion-path fallback
    sanity_gate: dict = {"passed": True, "metrics": []}
    for sg_candidate in [
        db_path.parent / "summary" / "sanity_gate.json",
        db_path.parent.parent / "sanity_gate.json",
    ]:
        if sg_candidate.exists():
            try:
                sg = json.loads(sg_candidate.read_text(encoding="utf-8"))
                sanity_gate["passed"] = bool(sg.get("passed", True))
                sanity_gate["metrics"] = [
                    {"name": k, "value": str(v), "threshold": "", "passed": True}
                    for k, v in sg.items() if k not in ("passed", "reason", "failures")
                ]
            except Exception:
                pass
            break

    return {
        "run_date":           datetime.now().strftime("%B %d, %Y %H:%M"),
        "db_name":            db_path.name,
        "total_records":      total,
        "actions":            actions,
        "n_with_changes":     int((df["fix_types"].fillna("") != "").sum()) if "fix_types" in df.columns else 0,
        "n_clean":            int((df["fix_types"].fillna("") == "").sum()) if "fix_types" in df.columns else 0,
        "sanity_gate":        sanity_gate,
        "match_sources":      match_sources,
        "confidence_bands":   confidence_bands,
        "active_zero_count":  active_zero_count,
        "active_zero_sample": active_zero_sample,
        "salary":             salary,
        "status_transitions": status_transitions,
        "hire_date_stats":    hire_date_stats,
        "wave_dates":         wave_dates,
        "reject_matches":     reject_matches,
        "review_queue":       review_queue,
        "findings":           findings,
        "change_summary":     change_summary,
    }


def _os_cleanup(path: str) -> None:
    """Remove a temp file, silently ignoring errors."""
    try:
        os.unlink(path)
    except Exception:
        pass


def _try_node_generator(df: pd.DataFrame, db_path: Path, wide_path: Path, out_path: Path) -> bool:
    """Attempt to generate the report using the Node.js/docx v9 generator.

    Returns True if the report was produced successfully.
    Returns False (caller falls back to python-docx) if:
      - node is not installed / not in PATH
      - the docx npm package is missing
      - build_report.js exits non-zero
      - the output file is absent or suspiciously small
    """
    # 1. Probe for node
    try:
        probe = subprocess.run(
            ["node", "--version"],
            capture_output=True, text=True, timeout=10,
        )
        if probe.returncode != 0:
            return False
        node_ver = probe.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False

    print(f"[generate_report] Node.js {node_ver} detected - trying JS generator")

    # 2. Locate build_report.js
    js_path = _HERE / "build_report.js"
    if not js_path.exists():
        print(f"[generate_report] build_report.js not found at {js_path} - falling back")
        return False

    # 3. Serialise run data to a temp JSON file
    try:
        data = _serialize_run_data(df, db_path, wide_path)
    except Exception as exc:
        print(f"[generate_report] serialize error: {exc} - falling back")
        return False

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".json")
    try:
        with os.fdopen(tmp_fd, "w", encoding="utf-8") as fh:
            json.dump(data, fh, ensure_ascii=False, default=str)

        # 4. Run node
        result = subprocess.run(
            ["node", str(js_path), "--data", tmp_path, "--out", str(out_path)],
            capture_output=True, text=True, timeout=120,
        )
        if result.stdout:
            print(result.stdout.rstrip())
        if result.stderr:
            print(result.stderr.rstrip(), file=sys.stderr)
        if result.returncode != 0:
            print(f"[generate_report] JS generator exited {result.returncode} - falling back")
            return False
        if not out_path.exists() or out_path.stat().st_size < 2048:
            print("[generate_report] JS generator output missing or too small - falling back")
            return False
        return True

    except subprocess.TimeoutExpired:
        print("[generate_report] JS generator timed out - falling back")
        return False
    except Exception as exc:
        print(f"[generate_report] JS generator error: {exc} - falling back")
        return False
    finally:
        _os_cleanup(tmp_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="Audit Report Generator.")
    parser.add_argument("--db",   default=None, metavar="PATH",
                        help=f"SQLite DB path (default: {DB_PATH}).")
    parser.add_argument("--wide", default=None, metavar="PATH",
                        help=f"wide_compare.csv path (default: {WIDE_CSV}).")
    parser.add_argument("--out",  default=None, metavar="PATH",
                        help=f"Output .docx path (default: {OUT_PATH}).")
    parser.add_argument("--no-node", action="store_true",
                        help="Skip Node.js generator and use python-docx directly.")
    args = parser.parse_args(argv)

    db_path   = Path(args.db)   if args.db   else DB_PATH
    wide_path = Path(args.wide) if args.wide else WIDE_CSV
    out_path  = Path(args.out)  if args.out  else OUT_PATH

    if not db_path.exists() and not wide_path.exists():
        print(f"[error] Neither DB ({db_path.name}) nor wide_compare.csv found.", file=sys.stderr)
        sys.exit(2)

    df = _load_data(db_path, wide_path)

    # Ensure numeric columns
    for col in ["salary_delta", "confidence", "priority_score"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    # ---------------------------------------------------------------------------
    # Try Node.js generator first (brand spec: navy headers, badges, callouts)
    # Falls back to python-docx automatically if node is unavailable.
    # ---------------------------------------------------------------------------
    if not args.no_node:
        if _try_node_generator(df, db_path, wide_path, out_path):
            try:
                display_path = out_path.relative_to(ROOT)
            except ValueError:
                display_path = out_path
            print(f"\n[generate_report] saved (JS/docx v9): {display_path}")
            return
        print("[generate_report] using python-docx fallback")

    # ---------------------------------------------------------------------------
    # Build document (python-docx fallback)
    # ---------------------------------------------------------------------------
    doc = Document()

    # Title page
    title   = doc.add_heading("Data Whisperer - Reconciliation Audit Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _BLUE_DARK

    subtitle   = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}  |  "
        f"Records: {len(df):,}  |  Source: {db_path.name}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Sections
    _section_executive_summary(doc, df)
    doc.add_page_break()
    _section_match_quality(doc, df)
    doc.add_page_break()
    _section_data_quality(doc, df)
    doc.add_page_break()
    _section_field_changes(doc, df)
    doc.add_page_break()
    _section_review_queue(doc, df)
    _section_rejected_matches(doc, df)
    doc.add_page_break()
    _section_recommendations(doc, df)

    # Save
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    try:
        display_path = out_path.relative_to(ROOT)
    except ValueError:
        display_path = out_path
    print(f"\n[generate_report] saved: {display_path}")
    print(f"  sections: Executive Summary, Match Quality, Data Quality,")
    print(f"            Field Changes, Review Queue, Rejected Matches, Recommendations")


if __name__ == "__main__":
    main()
